"""
Master Technical Architecture document generator.
Output: NotesApp_Architecture.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Styling helpers ----------

def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFBFBF", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 0:
            run.font.size = Pt(28)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    r = p.runs[0] if p.runs else p.add_run()
    if not p.runs:
        r = p.add_run(text)
    else:
        # Replace style placeholder text
        p.runs[0].text = text
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    if not p.runs:
        p.add_run(text)
    else:
        p.runs[0].text = text
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def add_callout(doc, label, text, fill="E8F1FA"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color="A8C5E0")
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    doc.add_paragraph()


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "1F3A5F")
        set_cell_borders(cell, color="1F3A5F", size="6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Body rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            if ri % 2 == 0:
                set_cell_shading(cell, "F4F7FB")
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Calibri'
            r.font.size = Pt(10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    set_cell_borders(cell, color="D9D9D9")
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    for line in code.split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
    doc.add_paragraph()


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------- Build the document ----------

doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ==================== TITLE PAGE ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(120)
r = title.add_run("Personal Notes & Knowledge Platform")
r.font.name = 'Calibri'
r.font.size = Pt(32)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Master Technical Architecture")
r.font.name = 'Calibri'
r.font.size = Pt(20)
r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run("An Evernote replacement with a fully integrated, fully local AI second brain. "
                 "V1 ships with Ollama only (running on the user's Legion desktop) for zero AI cost "
                 "and complete privacy. The AIProvider port architecture is preserved so Claude, "
                 "OpenAI, and other providers can be added later as single-file adapters. "
                 "Designed to scale from personal use to a public product.")
r.font.name = 'Calibri'
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, bold in [("Document Version: 1.0 (Draft for Review)", True),
                    ("Date: 2026-05-19", False),
                    ("Owner: V. Gakas", False),
                    ("Status: Pending review", False)]:
    r = meta.add_run(line + "\n")
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.bold = bold

page_break(doc)


# ==================== TABLE OF CONTENTS ====================
add_heading(doc, "Table of Contents", level=1)

toc_items = [
    ("1.  Executive Summary", "3"),
    ("2.  Vision, Scope & Success Criteria", "3"),
    ("3.  Functional Requirements", "4"),
    ("4.  Non-Functional Requirements", "5"),
    ("5.  High-Level Architecture", "6"),
    ("6.  Technology Stack", "8"),
    ("7.  Modular Component Breakdown (M1–M15)", "9"),
    ("8.  Data Model", "17"),
    ("9.  API Contract (Backend ↔ Clients)", "19"),
    ("10. Second-Brain AI Design (Multi-Provider)", "21"),
    ("11. Security, Privacy & Compliance", "24"),
    ("12. Phased Roadmap", "25"),
    ("13. Work-Package Distribution (Opus / Sonnet / Haiku)", "26"),
    ("14. Forward Compatibility & Multi-Tenant Path", "28"),
    ("15. Risks, Tradeoffs & Open Decisions", "29"),
    ("16. Appendix: Glossary & References", "30"),
    ("17. Ollama Deployment Guide (V1 Hardware-Specific)", "32"),
]
for text, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r = p.add_run(f"\t{page}")
    r.font.name = 'Calibri'
    r.font.size = Pt(11)

page_break(doc)


# ==================== 1. EXECUTIVE SUMMARY ====================
add_heading(doc, "1. Executive Summary", level=1)
add_para(doc,
    "This document defines the master technical architecture for a personal notes "
    "and knowledge-management platform intended to replace a paid Evernote subscription. "
    "The defining feature beyond Evernote is a fully integrated AI \"second brain\": an "
    "assistant that has indexed every note, every attachment's extracted text, every "
    "calendar event, and surfaces relevant context proactively across both clients.")

add_para(doc,
    "V1 ships with a single AI provider: Ollama, running on the user's home desktop "
    "(Lenovo Legion, Intel i7, RTX 3060 12GB VRAM, 16GB RAM). The phone and laptop reach "
    "this Ollama instance over Tailscale, so the second brain works from anywhere with no "
    "port-forwarding and no public exposure. Running cost for AI is effectively zero "
    "(local inference; electricity only). Privacy is maximal — note contents never leave "
    "the user's hardware.")

add_para(doc,
    "Critically, the AI is built behind a provider-agnostic AIProvider port. V1 only "
    "implements the OllamaAdapter, but adding Claude, OpenAI, Gemini, or any other "
    "provider later is a single adapter file plus a Settings entry — no app-layer changes. "
    "This is the architectural commitment that lets the user revisit the AI provider "
    "decision at any time without rebuilding anything.")

add_para(doc,
    "The system is initially scoped to a single user (the document owner), but every "
    "architectural decision below is constrained by a forward-compatibility requirement: "
    "the codebase must be ready to be repositioned as a public, multi-tenant SaaS product "
    "with minimal rework. This is achieved by treating every external dependency "
    "(storage provider, AI provider, calendar provider, auth provider, vector store) "
    "as a swappable port behind a stable interface.")

add_para(doc,
    "The deliverable spans two clients (an Android native app and a Next.js web app) "
    "backed by a single managed backend (Supabase: Postgres + Auth + Realtime + pgvector). "
    "User-owned files are stored in the user's own Google Drive or OneDrive — the platform "
    "holds metadata and AI-derived artefacts, not the source files. This keeps the privacy "
    "story clean and avoids storage cost as the system scales.")

add_callout(doc, "Key Principles",
    "(1) Discrete, swappable modules behind stable interfaces — any provider (storage, "
    "AI, calendar, vector DB) replaceable in under a day. (2) V1 = Ollama only; the "
    "multi-provider scaffolding exists so Claude, OpenAI, and others are a one-file add "
    "whenever the user wants them. (3) Privacy first — note bytes never leave the user's "
    "hardware in V1.")

page_break(doc)


# ==================== 2. VISION ====================
add_heading(doc, "2. Vision, Scope & Success Criteria", level=1)

add_heading(doc, "2.1 Vision", level=2)
add_para(doc,
    "Build a personal knowledge tool that does everything Evernote does, plus what "
    "Evernote never delivered: a true second brain. The AI knows your notes the way you "
    "wish you did — surfaces forgotten connections, drafts responses with your context, "
    "answers questions across years of notes in seconds. It travels with you on Android, "
    "is always one tap away, and respects that you (not the vendor) own both the data and "
    "the choice of which AI handles it.")

add_heading(doc, "2.2 In-Scope (V1)", level=2)
for item in [
    "Android mobile app (Kotlin + Jetpack Compose) — primary client.",
    "Web app (Next.js + React) — secondary client.",
    "Rich-text notes with attachments, notebooks, tags, and full-text search.",
    "Cross-device sync with offline-first behavior on Android.",
    "Second-brain AI powered by Ollama on the user's Legion desktop, reached via Tailscale.",
    "AI features: semantic search, conversational Q&A grounded in notes with citations, automatic tagging, summarization, daily briefing, connection discovery.",
    "Persistent AI conversations: chat history is stored locally; the AI remembers prior turns within a conversation.",
    "Per-task model routing within Ollama (e.g., llama3.1:8b for chat, llama3.2:3b for tagging).",
    "Two-way Google Calendar integration (events surface as notes; notes can create events).",
    "User-supplied cloud storage: Google Drive for file attachments (OneDrive adapter is V2).",
    "AIProvider port architecture, even though only one adapter ships — this is the modularity guarantee.",
]:
    add_bullet(doc, item)

add_heading(doc, "2.3 Out-of-Scope (V1) — architected for, but not shipped yet", level=2)
for item in [
    "Cloud AI providers (Claude, OpenAI, Gemini) — adapters exist in V2; port is ready for them.",
    "Voice capture — moved to V1.5; requires whisper.cpp setup on the Legion.",
    "OneDrive storage adapter — V2; only Google Drive in V1.",
    "iOS app — V3 or later.",
    "Web clipper browser extension.",
    "Real-time multi-user collaboration on the same note.",
    "End-to-end encryption with zero-knowledge keys (deferred to public launch).",
    "Self-hosted backend (V1 uses Supabase managed; migration path documented).",
]:
    add_bullet(doc, item)

add_heading(doc, "2.4 Success Criteria", level=2)
add_table(doc,
    headers=["#", "Criterion", "Measure"],
    rows=[
        ["1", "Feature parity with Evernote personal usage", "All recurring workflows can be performed in the new app within 1 month of cutover."],
        ["2", "Second brain delivers real value", "User asks the local AI a question about their corpus daily; chat with citations is the most-used feature after notes themselves."],
        ["3", "Local AI is fast enough to feel natural", "Chat with the second brain returns first token within 1.5s and streams at ≥ 25 tokens/sec on the Legion (RTX 3060)."],
        ["4", "Modularity proven", "Adding a Claude or OpenAI adapter in V2 requires only one new file in the codebase; no app-layer changes."],
        ["5", "Forward-compatible", "Adding a second user account requires no schema migration, only enabling RLS policies that are already in place."],
        ["6", "Operating cost", "Under $5/month for personal use (Supabase free tier, Tailscale free tier, AI on local hardware)."],
        ["7", "Privacy by default", "No note content leaves the user's hardware (Legion + Tailscale tunnel). Verifiable in network logs."],
    ],
    col_widths=[0.4, 2.6, 3.5])

page_break(doc)


# ==================== 3. FUNCTIONAL REQUIREMENTS ====================
add_heading(doc, "3. Functional Requirements", level=1)

add_heading(doc, "3.1 Notes & Editor", level=2)
for item in [
    "Create, edit, delete, and restore notes with rich text (headings, lists, checkboxes, code blocks, tables, inline images, links).",
    "Organize notes in a hierarchy of Notebooks (one level deep, like Evernote).",
    "Apply multiple tags per note. Tags are user-defined and AI-suggested.",
    "Pin a note. Star a note. Soft-delete with a 30-day trash retention.",
    "Attach files (images, PDFs, audio). Files live in the user's Drive/OneDrive; only references are stored in the platform DB.",
    "Markdown view and toggle. Internal storage format is ProseMirror JSON (lossless to/from Markdown).",
]:
    add_bullet(doc, item)

add_heading(doc, "3.2 Search & Discovery", level=2)
for item in [
    "Full-text keyword search across title and body.",
    "Tag and notebook filters that compose with text search.",
    "Semantic search: 'find notes about X', returning ranked results by embedding similarity.",
    "Recent and recently-modified views.",
    "Saved searches (filter combinations the user reuses).",
]:
    add_bullet(doc, item)

add_heading(doc, "3.3 AI Capabilities — The Second Brain", level=2)

add_para(doc, "Per-note actions", bold=True)
for item in [
    "Summarize at three lengths (one-line, paragraph, bullet outline).",
    "Extract action items as a checklist; optional 'create reminders' affordance.",
    "Suggest 3–5 tags; user accepts/rejects with one tap.",
    "Suggest a title; auto-applied on first save if title field is empty.",
    "Improve writing / rewrite tone (slash command in the editor).",
]:
    add_bullet(doc, item)

add_para(doc, "Corpus-wide actions", bold=True)
for item in [
    "Chat with your notes: free-form Q&A grounded in your corpus, with clickable citations to source notes.",
    "Persistent conversations: history is stored; the AI remembers context across sessions; user can start a fresh thread any time.",
    "Connection discovery: open any note and tap 'related notes' to see semantically-linked older notes.",
    "Daily briefing: optional morning summary — what you wrote yesterday, what's due today, open action items, calendar overview.",
    "Voice capture (Android): hold a button, dictate; the AI transcribes, auto-titles, auto-tags, files into the right notebook.",
    "Quick-ask from anywhere on Android: a persistent notification / quick-tile opens a chat without launching the full app.",
]:
    add_bullet(doc, item)

add_para(doc, "Provider control (V1 = single provider, multi-provider ready)", bold=True)
for item in [
    "V1 ships with Ollama only. The AIProvider port supports adding more providers; the Settings UI starts with a single 'Provider: Ollama' row.",
    "Settings → AI lets the user pick the model per task within Ollama (e.g., llama3.1:8b for chat, llama3.2:3b for tagging, nomic-embed-text for embeddings).",
    "Ollama endpoint URL is configurable (default points to the user's Tailscale-reachable Legion).",
    "When V2 ships additional providers, switching is non-destructive: notes, conversations, and embeddings are preserved; the new provider just takes over future requests.",
    "Streaming responses for chat (token-by-token rendering on both clients).",
    "AI runs only on explicit user action (or on background indexing of new notes for embeddings) — no silent corpus scraping.",
]:
    add_bullet(doc, item)

add_heading(doc, "3.4 Sync & Offline (Android)", level=2)
for item in [
    "Read all cached notes offline.",
    "Create and edit notes offline; sync resolves on next connectivity.",
    "Conflict resolution: last-writer-wins on body, union on tags, with a 'view conflict' affordance.",
    "Realtime updates when both clients are online (Supabase Realtime channels).",
]:
    add_bullet(doc, item)

add_heading(doc, "3.5 Google Calendar Integration", level=2)
for item in [
    "OAuth-link the user's Google account once.",
    "Pull events for a configurable date range; events are read-only mirrored as a special note type.",
    "From any note, create a Calendar event (date, time, reminder, attendees).",
    "Daily note view shows today's calendar items inline with notes due today.",
]:
    add_bullet(doc, item)

add_heading(doc, "3.6 Storage Integration (Drive / OneDrive)", level=2)
for item in [
    "User picks one provider at setup (configurable; provider is recorded per user, not global).",
    "On attachment upload, the file is uploaded to the user's chosen provider; a reference (provider, file ID, name, mime, size) is stored in our DB.",
    "On attachment open, the client fetches a short-lived signed URL via the provider's SDK.",
    "The platform never copies attachment bytes into its own storage; the only platform-resident binary is the embedding vector for AI search.",
]:
    add_bullet(doc, item)

page_break(doc)


# ==================== 4. NON-FUNCTIONAL ====================
add_heading(doc, "4. Non-Functional Requirements", level=1)

add_table(doc,
    headers=["Category", "Requirement", "Target"],
    rows=[
        ["Modularity", "Every external provider is behind an interface; no app code imports a provider SDK directly.", "Mandatory"],
        ["Forward compatibility", "Schema and code support multi-user from day one (RLS enforced even with one user).", "Mandatory"],
        ["Performance — Android", "Cold start to note list visible", "< 1.5s on mid-tier device"],
        ["Performance — Search", "Keyword search response", "< 200ms for 10k notes"],
        ["Performance — Semantic search", "End-to-end (embedding + retrieval + answer)", "< 4s p95"],
        ["Reliability", "Sync conflicts must not silently drop user content", "0 data loss"],
        ["Cost (personal)", "Total monthly OpEx", "≤ $25/month"],
        ["Cost (public, per active user)", "Marginal OpEx", "≤ $0.50/user/month at 1k users"],
        ["Security", "All data encrypted at rest (Supabase default) and in transit (TLS 1.2+).", "Mandatory"],
        ["Privacy", "User's attachment bytes never touch our infra.", "Mandatory"],
        ["Portability", "User can export full corpus (notes + metadata) as a single .zip of Markdown + JSON.", "Mandatory"],
        ["Observability", "Structured logs, error reporting (Sentry), basic usage metrics.", "Mandatory"],
    ],
    col_widths=[1.4, 3.6, 1.5])

page_break(doc)


# ==================== 5. HIGH-LEVEL ARCHITECTURE ====================
add_heading(doc, "5. High-Level Architecture", level=1)

add_para(doc,
    "The system uses a clean hexagonal (ports-and-adapters) architecture. The "
    "application core knows about domain concepts (Note, Notebook, Tag, Attachment, "
    "Event, Embedding) and depends only on a set of port interfaces. Adapters implement "
    "those ports against concrete providers (Supabase, Google Drive, OneDrive, Claude, "
    "Google Calendar). The adapter set is configured at startup, not hard-coded.")

add_heading(doc, "5.1 Layered View", level=2)
add_code_block(doc, """
┌───────────────────────────────────────────────────────────────────┐
│  CLIENTS                                                          │
│  ┌──────────────────────┐    ┌──────────────────────────────┐    │
│  │ Android (Kotlin +    │    │ Web (Next.js + React +       │    │
│  │ Jetpack Compose +    │    │ Tiptap editor + shadcn/ui)   │    │
│  │ Room offline cache)  │    │                              │    │
│  └──────────┬───────────┘    └──────────────┬───────────────┘    │
└─────────────┼────────────────────────────────┼───────────────────┘
              │            HTTPS / WSS         │
┌─────────────┴────────────────────────────────┴───────────────────┐
│  BACKEND (Supabase project + Edge Functions)                     │
│  ┌─────────────────────────────────────────────────────────┐     │
│  │ Auth (Supabase Auth + Google OAuth)                     │     │
│  ├─────────────────────────────────────────────────────────┤     │
│  │ Postgres (notes, notebooks, tags, attachments_refs,     │     │
│  │           events_mirror, embeddings via pgvector)       │     │
│  │ Row-Level Security enforced on every table              │     │
│  ├─────────────────────────────────────────────────────────┤     │
│  │ Realtime (sync channels per user)                       │     │
│  ├─────────────────────────────────────────────────────────┤     │
│  │ Edge Functions (Deno): AI orchestration, OAuth token    │     │
│  │ refresh, Drive/OneDrive proxy when needed               │     │
│  └─────────────────────────────────────────────────────────┘     │
└────────────┬────────────────┬───────────────┬────────────────────┘
             │                │               │
   ┌─────────┴───┐   ┌────────┴────┐   ┌──────┴──────────┐
   │ Claude API  │   │ Google APIs │   │ Microsoft Graph │
   │ (Anthropic) │   │ Drive +     │   │ (OneDrive)      │
   │             │   │ Calendar    │   │                 │
   └─────────────┘   └─────────────┘   └─────────────────┘
""")

add_heading(doc, "5.2 Ports & Adapters", level=2)
add_para(doc, "Every external dependency is fronted by a port interface. Concrete "
              "adapters are injected at composition root. This is the heart of the "
              "modularity requirement.")
add_table(doc,
    headers=["Port (Interface)", "V1 Adapter", "Future Swap Candidates"],
    rows=[
        ["StorageProvider", "GoogleDriveAdapter", "OneDriveAdapter, S3Adapter, DropboxAdapter, LocalFsAdapter"],
        ["AIProvider", "ClaudeAdapter, OpenAIAdapter, OllamaAdapter (all V1)", "GeminiAdapter, AzureOpenAIAdapter, GroqAdapter, MistralAdapter"],
        ["EmbeddingProvider", "VoyageAdapter (or ClaudeViaProxy)", "OpenAIEmbeddings, sentence-transformers (self-hosted)"],
        ["VectorStore", "PgVectorAdapter", "PineconeAdapter, WeaviateAdapter, QdrantAdapter"],
        ["CalendarProvider", "GoogleCalendarAdapter", "OutlookCalendarAdapter, CalDAVAdapter"],
        ["AuthProvider", "SupabaseAuthAdapter (Google OAuth)", "Auth0Adapter, ClerkAdapter, CognitoAdapter"],
        ["NotificationProvider", "FCMAdapter (Android)", "APNSAdapter (iOS), WebPushAdapter"],
    ],
    col_widths=[1.8, 2.0, 2.7])

add_callout(doc, "Modularity Test",
    "If a port can be re-implemented without grep-ing for the old provider's name "
    "anywhere outside the adapter file, the boundary is correct.")

page_break(doc)


# ==================== 6. TECH STACK ====================
add_heading(doc, "6. Technology Stack", level=1)

add_table(doc,
    headers=["Layer", "Choice", "Reason"],
    rows=[
        ["Android app", "Kotlin + Jetpack Compose", "Native UX, Google's modern stack, future-proof, best long-term."],
        ["Android local store", "Room (SQLite) + DataStore", "Mature offline cache; integrates with Compose flows."],
        ["Android sync", "WorkManager + Supabase Realtime", "Battery-aware, survives process death."],
        ["Web app", "Next.js 14 (App Router) + React 18", "Server components where useful; familiar; good Supabase integration."],
        ["Web editor", "Tiptap (ProseMirror)", "Best-in-class rich-text editor with extensibility for AI inline actions."],
        ["Web UI kit", "shadcn/ui + Tailwind", "Owned components, easy theming for future white-label."],
        ["Backend", "Supabase (managed)", "Postgres + Auth + Realtime + Storage + pgvector in one. Self-host path exists."],
        ["Vector search", "pgvector inside Supabase Postgres", "No second datastore; HNSW indexes fast enough at single-user scale."],
        ["Server logic", "Supabase Edge Functions (Deno + TypeScript)", "Co-located with DB; cheap; easy to migrate to a standalone Node service."],
        ["AI provider (V1)", "Ollama on user's Legion desktop (Windows 11, i7, RTX 3060 12GB, 16GB RAM)", "Local inference; zero per-query cost; data never leaves user hardware."],
        ["Reachability", "Tailscale (mesh VPN, free tier)", "Phone/laptop reach Ollama from anywhere; no port-forwarding; no public exposure."],
        ["Chat / Q&A model", "llama3.1:8b (default) or qwen2.5:7b (alt)", "Fits in 12GB VRAM; 30–60 tok/sec on RTX 3060."],
        ["Bulk tasks (tag/title)", "llama3.2:3b", "Smaller, 2–3× faster; quality sufficient for tagging."],
        ["Embeddings", "nomic-embed-text via Ollama (768-dim)", "Free, local, standard for Ollama. Vectors namespaced in DB so a future switch is non-destructive."],
        ["Voice (V1.5, optional)", "whisper.cpp on the Legion or OpenAI Whisper API", "User chooses in Settings; default skipped in V1."],
        ["Future providers (V2)", "Claude, OpenAI, Gemini — one adapter file each", "AIProvider port designed for this from day one."],
        ["File storage", "User's Google Drive or OneDrive", "Zero-cost to platform; user owns bytes; privacy-clean."],
        ["Calendar", "Google Calendar API v3", "User's existing calendar. Outlook adapter planned."],
        ["Auth", "Supabase Auth with Google OAuth", "Single sign-on with Google (which is also our Calendar/Drive grant)."],
        ["Observability", "Sentry (errors), PostHog (product analytics)", "Free tiers cover personal scale."],
        ["CI/CD", "GitHub Actions", "Standard. Triggers deploy to Vercel (web) and Play Store internal track (Android)."],
    ],
    col_widths=[1.4, 2.2, 3.0])

add_callout(doc, "Note on AI provider strategy",
    "V1 ships one adapter (OllamaAdapter) but keeps the full AIProvider port abstraction. "
    "This is intentional: the user gets free, private AI immediately while the codebase remains "
    "ready to accept Claude, OpenAI, Gemini, Groq, Mistral or Azure OpenAI as drop-in adapters. "
    "Each future provider = one new file. The architecture proves itself by NOT requiring a "
    "second adapter in V1 to remain swappable.")

page_break(doc)


# ==================== 7. MODULES ====================
add_heading(doc, "7. Modular Component Breakdown", level=1)
add_para(doc,
    "Each subsection below is a self-contained work package. It can be assigned to a "
    "developer (human or AI) with no further context beyond the interface contract and "
    "this document. Modules are intentionally numbered for cross-reference; module IDs "
    "(M1–M14) appear in Section 13 (work-package distribution).")


def module_block(doc, mid, title, purpose, interface, responsibilities, deliverables, depends_on, suggested_developer):
    add_heading(doc, f"7.{mid[1:]} {mid} — {title}", level=2)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, purpose)

    add_para(doc, "Public interface (port)", bold=True)
    add_code_block(doc, interface)

    add_para(doc, "Responsibilities", bold=True)
    for r in responsibilities:
        add_bullet(doc, r)

    add_para(doc, "Deliverables", bold=True)
    for d in deliverables:
        add_bullet(doc, d)

    add_para(doc, "Depends on", bold=True)
    add_para(doc, depends_on)

    add_para(doc, "Suggested developer", bold=True, italic=False)
    add_para(doc, suggested_developer, italic=True)
    doc.add_paragraph()


module_block(doc, "M1", "Domain Model & Shared Types",
    "Define the canonical domain entities — Note, Notebook, Tag, Attachment, "
    "Event, Embedding, User — as plain data classes with no provider knowledge. "
    "This module is the lingua franca every other module imports.",
    """// Kotlin (shared via KMP later if desired; for now duplicated in TS for web)
data class Note(
  val id: NoteId,
  val ownerId: UserId,
  val notebookId: NotebookId?,
  val title: String,
  val bodyJson: String,            // ProseMirror JSON
  val tags: Set<TagId>,
  val createdAt: Instant,
  val updatedAt: Instant,
  val isPinned: Boolean,
  val isTrashed: Boolean,
  val attachmentRefs: List<AttachmentRef>
)

data class AttachmentRef(
  val id: AttachmentId,
  val provider: StorageProviderId, // GOOGLE_DRIVE | ONEDRIVE | ...
  val externalFileId: String,
  val mimeType: String,
  val sizeBytes: Long,
  val displayName: String
)""",
    [
        "Pure data classes / TypeScript types — no I/O.",
        "ID types are opaque (value classes / branded types) to prevent mix-ups.",
        "Validation invariants are enforced in constructors.",
    ],
    [
        "Kotlin module `core-domain` (Android + future KMP).",
        "TypeScript package `@app/domain` (web + edge functions).",
        "Unit tests for invariants.",
    ],
    "Nothing.",
    "Sonnet — well-specified, mechanical work.")


module_block(doc, "M2", "Auth Module",
    "Handle sign-in (Google OAuth via Supabase), token storage on Android, session "
    "lifecycle on web. Expose a stable `AuthProvider` port; Supabase is the V1 adapter.",
    """interface AuthProvider {
  suspend fun signInWithGoogle(): AuthResult
  suspend fun signOut()
  fun currentUser(): Flow<User?>
  suspend fun accessToken(): String?       // for backend calls
  suspend fun providerAccessToken(provider: ExternalProviderId): String?
  //                                       ^ Drive / Calendar etc.
}""",
    [
        "Initiate OAuth flow on Android (Chrome Custom Tabs) and web (Supabase JS).",
        "Persist session securely (EncryptedSharedPreferences on Android; httpOnly cookie on web).",
        "Refresh provider access tokens (Drive, Calendar) automatically when expired.",
        "Expose currentUser as a reactive Flow / hook.",
    ],
    [
        "Android: `auth-android` library + composable `SignInScreen`.",
        "Web: `useAuth` hook + `<AuthGate>` component.",
        "Edge function: `auth/refresh-provider-token`.",
    ],
    "M1.",
    "Sonnet for adapters; Opus to design the token-refresh strategy across providers.")


module_block(doc, "M3", "Storage Provider Module",
    "Abstract file storage. V1 ships Google Drive and OneDrive adapters; both implement "
    "the same `StorageProvider` port. User picks one at setup; switching is supported "
    "but does not migrate existing attachments (that is a V2 feature).",
    """interface StorageProvider {
  val id: StorageProviderId
  suspend fun upload(file: FileInput, folderHint: String?): RemoteFile
  suspend fun getDownloadUrl(externalId: String, ttlSeconds: Int = 300): String
  suspend fun delete(externalId: String)
  suspend fun ensureAppFolder(): String     // returns root folder ID
}""",
    [
        "GoogleDriveAdapter: uses Drive v3 REST API with the user's OAuth token.",
        "OneDriveAdapter: uses Microsoft Graph API with MSAL.",
        "Both create and reuse a single app-scoped folder (`/NotesApp/`) so the user can audit what we wrote.",
        "Upload returns a `RemoteFile` { externalId, name, mime, size }.",
        "Signed/download URLs are short-lived (default 5 min) and never persisted.",
    ],
    [
        "Library `storage-google-drive`.",
        "Library `storage-onedrive`.",
        "Integration tests against test accounts.",
        "ADR documenting the chosen quota / scope model.",
    ],
    "M1, M2 (needs provider access tokens).",
    "Opus — OAuth scopes and quota behaviors differ between providers and need careful design.")


module_block(doc, "M4", "Notes Core Service",
    "CRUD for notes, notebooks, tags. Pure backend logic — does not know about Drive, "
    "Claude, or Calendar. Lives in Supabase (Postgres tables + RLS + a few Edge Functions "
    "for compound operations).",
    """interface NotesService {
  suspend fun createNote(input: NewNoteInput): Note
  suspend fun updateNote(id: NoteId, patch: NotePatch): Note
  suspend fun trashNote(id: NoteId)
  suspend fun restoreNote(id: NoteId)
  suspend fun listNotes(filter: NoteFilter, page: Page): PagedResult<Note>
  suspend fun getNote(id: NoteId): Note
  suspend fun searchKeyword(q: String, filter: NoteFilter): List<NoteHit>
}""",
    [
        "Schema (Postgres) for notes, notebooks, tags, note_tags, attachments_refs.",
        "Row-Level Security policies enforce owner == auth.uid() on every table.",
        "Postgres full-text search index (tsvector) for keyword search.",
        "Triggers update `updated_at` on every change.",
        "30-day trash retention via a scheduled SQL job (pg_cron).",
    ],
    [
        "SQL migration files (versioned with supabase/migrations).",
        "TypeScript client SDK auto-generated from PostgREST types.",
        "Edge function `notes/bulk-update` for atomic multi-note operations.",
    ],
    "M1.",
    "Sonnet — schema + RLS + standard CRUD endpoints.")


module_block(doc, "M5", "Sync Engine (Android)",
    "Offline-first sync between Android Room cache and the Supabase backend. Handles "
    "queuing of offline writes, conflict detection, and realtime push.",
    """interface SyncEngine {
  fun start(scope: CoroutineScope)
  fun syncStatus(): StateFlow<SyncStatus>
  suspend fun forcePull()
  suspend fun forcePush()
}""",
    [
        "Outbox table in Room for pending writes; periodic flush via WorkManager.",
        "Pull strategy: server `updated_at > last_sync_at` per table, paged.",
        "Push strategy: outbox row → REST call → on success, remove from outbox.",
        "Conflict policy: last-writer-wins on body; union on tags; flag conflicts in a system note for user awareness.",
        "Subscribe to Supabase Realtime channel for the user's row updates while app is foregrounded.",
    ],
    [
        "Kotlin library `sync-android`.",
        "Coroutines-based; testable with a fake `NotesService`.",
        "Doc: 'How sync works' explainer for future contributors.",
    ],
    "M1, M4.",
    "Opus — concurrency, idempotency, and conflict policy are the hardest part of the system.")


module_block(doc, "M6", "Editor Module (Web + Android)",
    "Rich-text editor with a shared JSON document model (ProseMirror schema) so notes "
    "render identically on both clients.",
    """// Document model: ProseMirror JSON
// Web: Tiptap editor with custom extensions
// Android: a Compose renderer + editor that reads/writes the same JSON

interface NoteEditor {
  fun load(json: String)
  fun toJson(): String
  fun toMarkdown(): String
  fun insertAttachment(ref: AttachmentRef)
  fun registerAICommand(cmd: AICommand)   // e.g., 'Summarize selection'
}""",
    [
        "Define a frozen ProseMirror schema (headings, paragraphs, lists, checkboxes, code, image, attachment, internal-link).",
        "Web: implement with Tiptap; ship custom nodes for attachment and internal-link.",
        "Android: render in Compose; for V1, edit in a simplified subset and full WYSIWYG via a CommonMark view.",
        "Inline AI actions: 'Summarize', 'Improve writing', 'Extract action items' as slash-commands.",
    ],
    [
        "Web package `editor-web`.",
        "Android module `editor-android`.",
        "JSON ↔ Markdown converter (shared logic, mirrored).",
    ],
    "M1, M3 (for attachment insert), M9 (for AI commands).",
    "Opus on Android (Compose rich-text is non-trivial); Sonnet on web (Tiptap is well-trodden).")


module_block(doc, "M7", "Attachment Pipeline",
    "Bridge between the editor (which knows about an Attachment node) and the Storage "
    "Provider (which doesn't know what a note is). Encapsulates upload progress, retry, "
    "thumbnail generation, and reference persistence.",
    """interface AttachmentPipeline {
  suspend fun attach(noteId: NoteId, file: FileInput): AttachmentRef
  suspend fun resolveUrl(ref: AttachmentRef): String
  suspend fun remove(ref: AttachmentRef)
  fun uploadProgress(ref: AttachmentRef): Flow<Progress>
}""",
    [
        "Calls StorageProvider.upload() (M3), gets externalId, writes AttachmentRef to DB (M4).",
        "Generates thumbnails for images on the client before upload (saves bandwidth and Drive quota).",
        "Handles 'attachment dangling': scheduled job warns if a note references a missing external file.",
    ],
    [
        "Android library `attachments-android`.",
        "Web library `attachments-web`.",
        "Edge function `attachments/validate` (periodic dangling-ref check).",
    ],
    "M3, M4.",
    "Sonnet.")


module_block(doc, "M8", "Calendar Integration Module",
    "Two-way bridge to Google Calendar. V1 ships Google; OutlookCalendarAdapter is a "
    "later swap behind the same port.",
    """interface CalendarProvider {
  val id: CalendarProviderId
  suspend fun listEvents(range: DateRange): List<CalendarEvent>
  suspend fun createEvent(input: NewEventInput): CalendarEvent
  suspend fun updateEvent(id: ExternalEventId, patch: EventPatch): CalendarEvent
  suspend fun deleteEvent(id: ExternalEventId)
  fun watchEvents(range: DateRange): Flow<CalendarChange>  // poll or push
}""",
    [
        "Use Google Calendar API v3 with the user's OAuth token (granted at sign-in).",
        "Mirror events into a Postgres table `events_mirror` for offline read and unified search.",
        "Outbound: from a note, 'Create Calendar event' opens a dialog; on submit, calls createEvent and stores back-reference.",
        "Polling cadence: every 15 min while foregrounded; webhook channels (Calendar push notifications) added in V2.",
    ],
    [
        "Edge function `calendar/sync`.",
        "Android UI: 'Today' view that lists events + notes due today.",
        "Web UI: same Today view.",
    ],
    "M1, M2.",
    "Sonnet.")


module_block(doc, "M9", "Second-Brain AI Services (Multi-Provider)",
    "The heart of the product. All AI capabilities — summarization, tagging, "
    "Q&A, chat-with-notes, daily briefing, connection discovery, voice capture — sit "
    "behind a provider-agnostic AIServices facade. Underneath, the AIProvider port has "
    "three V1 adapters (Claude, OpenAI, Ollama). User selects the active provider per "
    "task in Settings (see M15). This module is what makes the product a 'second brain' "
    "rather than 'notes with a chat button.'",
    """// Provider-agnostic port (all adapters implement this)
interface AIProvider {
  val id: ProviderId                            // CLAUDE | OPENAI | OLLAMA | ...
  val capabilities: Set<Capability>             // STREAMING, EMBEDDINGS, VISION, TOOL_USE, JSON_MODE
  val models: List<ModelDescriptor>             // available models for this provider
  suspend fun chat(messages: List<Message>, opts: AIOptions): AIResponse
  fun streamChat(messages: List<Message>, opts: AIOptions): Flow<AIChunk>
  suspend fun embed(texts: List<String>, model: String): List<FloatArray>?  // null if unsupported
  suspend fun ping(): ProviderHealth            // for the Settings 'Test connection' button
}

// V1 adapters: ClaudeAdapter, OpenAIAdapter, OllamaAdapter

// Domain-facing facade — clients only ever talk to this
interface AIServices {
  // Per-note (one-shot)
  suspend fun summarize(noteId: NoteId, length: SummaryLength): String
  suspend fun suggestTags(noteId: NoteId): List<String>
  suspend fun suggestTitle(noteId: NoteId): String
  suspend fun extractActionItems(noteId: NoteId): List<String>
  suspend fun rewrite(text: String, instruction: String): String

  // Corpus-wide (the second-brain features)
  fun chat(conversationId: ConversationId?, message: String, scope: NoteScope): Flow<ChatChunk>
  suspend fun relatedNotes(noteId: NoteId, k: Int = 10): List<RelatedNote>
  suspend fun dailyBriefing(date: LocalDate): Briefing
  suspend fun captureFromVoice(audio: ByteArray): VoiceCapture  // transcribe + classify + file

  // Conversation persistence (so the AI remembers)
  suspend fun listConversations(): List<Conversation>
  suspend fun getConversation(id: ConversationId): Conversation
  suspend fun newConversation(seed: String? = null): ConversationId
  suspend fun deleteConversation(id: ConversationId)
}""",
    [
        "AIServices is the only thing UI talks to. It resolves to the configured Ollama model per task via M15.",
        "OllamaAdapter — calls the user's Ollama endpoint URL (V1: a Tailscale hostname pointing at the user's Legion). Supports any model the user has pulled via `ollama pull`.",
        "Adapter declares capabilities at runtime by calling `GET /api/tags` to enumerate installed models. Settings UI populates model dropdowns from this list.",
        "RAG pipeline for chat: vector retrieval (M10) → top-k chunks → prompt assembly → stream answer with [[NoteId:xxx]] citations the UI converts to deep links.",
        "Persistent conversations stored in Postgres (`ai_conversations`, `ai_messages`); UI shows a history sidebar like ChatGPT.",
        "Daily briefing: scheduled per-user at user-configured time (default 8am); reads yesterday's notes + today's calendar + open action items; produces a Markdown summary delivered as an in-app notification.",
        "Connection discovery: 'related notes' uses pure vector similarity (no LLM call) — instant; tappable in editor.",
        "Voice capture (V1.5): audio uploaded to backend, transcribed by whisper.cpp running alongside Ollama on the Legion; AI classifies into existing notebook + suggests tags; user confirms or edits before save.",
        "All AI calls go through backend Edge Functions, which connect to Ollama over the user's Tailscale tunnel. The Legion is never directly exposed to the internet.",
        "Future adapters (Claude, OpenAI, etc.) implement the same AIProvider port. The day they're added, M15's routing dropdowns gain a Provider column and per-task overrides become available.",
        "Cost telemetry per task in `ai_usage_log` — in V1 this records latency and tokens but $0 cost.",
    ],
    [
        "Edge functions: `ai/chat` (SSE), `ai/summarize`, `ai/suggest-tags`, `ai/suggest-title`, `ai/related`, `ai/briefing`.",
        "One provider adapter (OllamaAdapter) with integration tests against a real Ollama instance.",
        "Prompt library in repo (`/prompts/*.md`) — provider-neutral templates; per-model overrides ready for V2.",
        "Conversation persistence schema (see §8).",
        "Voice capture pipeline (V1.5 — see §17.6).",
    ],
    "M1, M4 (reads notes), M10 (vector retrieval), M15 (routing config), §17 (Ollama deployment).",
    "Opus — even with a single adapter, RAG quality, conversation summarization for long history, and prompt design for open-weight models (which need more careful prompting than Claude/GPT) all require design judgment. The Ollama adapter itself is Sonnet-level.")


module_block(doc, "M10", "Embedding & Vector Search Module",
    "Manage embeddings for every note. Re-embed on update. Provide top-k retrieval "
    "for RAG. V1 uses pgvector inside Supabase Postgres; the `VectorStore` port supports "
    "swap to Pinecone or Qdrant later without retrieval-code changes.",
    """interface VectorStore {
  suspend fun upsert(noteId: NoteId, chunks: List<EmbeddingChunk>)
  suspend fun delete(noteId: NoteId)
  suspend fun search(queryEmbedding: FloatArray, k: Int, filter: VectorFilter): List<VectorHit>
}""",
    [
        "Chunk note body into ~500-token windows with overlap before embedding.",
        "Trigger: on note insert/update, an Edge Function enqueues an embedding job.",
        "pgvector index: HNSW with cosine distance.",
        "Filter pushdown: VectorFilter supports notebook, tag, date range — translated to SQL WHERE clauses (essential for multi-tenant).",
    ],
    [
        "Migration adding `note_embeddings` table.",
        "Edge function `embeddings/index` (idempotent, retryable).",
        "Backfill script that embeds the entire corpus on first run.",
    ],
    "M1, M4.",
    "Sonnet — well-bounded; the only judgment call is chunking strategy (document in code).")


module_block(doc, "M11", "Notification Module",
    "Push notifications for reminders and (later) shared-note events. V1 only sends "
    "local reminders for note-attached due dates; V2 adds cross-device push via FCM.",
    """interface NotificationProvider {
  suspend fun scheduleLocal(reminder: Reminder)
  suspend fun cancelLocal(reminderId: ReminderId)
  suspend fun sendPush(userId: UserId, payload: PushPayload)   // V2
}""",
    [
        "Android: AlarmManager + WorkManager for local reminders.",
        "Web: Web Notifications API for in-tab reminders.",
        "V2: Firebase Cloud Messaging adapter behind the same port.",
    ],
    [
        "Android library `notifications-android`.",
        "Web hook `useNotifications`.",
    ],
    "M1.",
    "Haiku — small, mechanical, well-documented platform APIs.")


module_block(doc, "M12", "Android App Shell",
    "Top-level Android app: navigation graph, theme, dependency-injection wiring "
    "(Hilt), splash, sign-in, main shell. This is the composition root that wires "
    "every module together for the Android client.",
    """// Composition root
@HiltAndroidApp
class NotesApp : Application() {
  // Modules provide: AuthProvider, NotesService client,
  // StorageProvider (chosen per user), CalendarProvider,
  // AIServices (proxies edge functions), SyncEngine,
  // NotificationProvider
}""",
    [
        "Navigation: Compose Navigation, single-activity.",
        "Theming: Material 3 with dynamic color; dark mode default.",
        "DI: Hilt; one module per port so swaps are mechanical.",
        "Crash reporting: Sentry SDK initialized in onCreate.",
    ],
    [
        "App module compiles and runs end-to-end with sign-in → note list → editor.",
        "Internal-track Play Store build.",
    ],
    "M1–M11.",
    "Opus for the initial scaffold (DI graph + navigation); Sonnet for screen implementations.")


module_block(doc, "M13", "Web App Shell",
    "Next.js application that wires the same set of ports to web-appropriate adapters. "
    "Reuses the TypeScript domain package from M1.",
    """// app/layout.tsx is the composition root for providers:
// AuthProvider, NotesService (Supabase JS client),
// StorageProvider (chosen per user),
// CalendarProvider, AIServices (fetch to edge functions),
// NotificationProvider""",
    [
        "Next.js 14 App Router; React Server Components for note list pages.",
        "Tiptap editor with the custom schema from M6.",
        "Streaming chat UI for M9's chat-with-notes (SSE).",
        "Responsive layout — tablet-first; phone web is acceptable but Android is the primary mobile target.",
    ],
    [
        "Web app deployable to Vercel with one env file.",
        "Lighthouse score ≥ 90 on the main editor page.",
    ],
    "M1–M11.",
    "Sonnet — modern Next.js work is well-trodden; Opus only if AI streaming UI needs design.")


module_block(doc, "M15", "AI Settings, Routing & Endpoint Management",
    "The user-facing control plane for the second brain. In V1, lets the user configure "
    "the Ollama endpoint URL, pick which model handles which task, test the connection, "
    "and see latency/usage telemetry. The routing schema is multi-provider-aware from "
    "day one so V2's added providers slot in without redesign.",
    """interface AIRouting {
  suspend fun providersAvailable(): List<ProviderId>             // configured AND reachable
  suspend fun getRouting(): RoutingConfig
  suspend fun setRouting(config: RoutingConfig)
  suspend fun resolveFor(task: AITask): AdapterBinding           // used by M9 internally
}

// Per-user routing config (stored in users_profile.ai_prefs jsonb)
data class RoutingConfig(
  val chat:       AdapterBinding,                 // e.g., {claude, claude-sonnet-4-6}
  val summarize:  AdapterBinding,                 // e.g., {claude, claude-haiku-4-5}
  val tagging:    AdapterBinding,                 // e.g., {ollama, llama3:8b}
  val titling:    AdapterBinding,                 // e.g., {openai, gpt-4o-mini}
  val embeddings: AdapterBinding,                 // e.g., {voyage, voyage-3-large}
  val briefing:   AdapterBinding,
  val voice:      AdapterBinding,
  val rewrite:    AdapterBinding,
)

interface AIKeyStore {
  suspend fun putKey(provider: ProviderId, secret: String)
  suspend fun hasKey(provider: ProviderId): Boolean
  suspend fun deleteKey(provider: ProviderId)
  // keys are encrypted at rest (Supabase Vault) and never returned to clients
  suspend fun ollamaEndpoint(): String?
  suspend fun setOllamaEndpoint(url: String)
}""",
    [
        "Settings screen on Android (Compose) and Web (Next.js) with three sections: Endpoint, Routing, Usage.",
        "Endpoint section (V1): single 'Ollama endpoint URL' input (default suggestion: the user's Tailscale hostname, e.g. http://legion.tailnet.ts.net:11434). Tap 'Test connection' to verify and enumerate installed models.",
        "Routing section: a table — left column is task (Chat / Summarize / Tag / Title / Embed / Briefing / Voice / Rewrite); right column in V1 is a single Model dropdown populated from the Ollama endpoint. In V2 (multi-provider) the same UI gains a Provider column.",
        "Usage section: chart of tokens used and latency per task per day; CSV export. Cost column reads $0 in V1 (local inference) but is wired up for V2.",
        "Endpoint URL stored in `users_profile.ai_prefs` (no API key needed for Ollama in V1).",
        "V2 readiness: ai_keys table and Supabase Vault integration are built in V1 even though unused, so adding Claude/OpenAI in V2 is a UI-only change to expose the inputs.",
        "Capability-aware UI: dropdowns only show models that support the task. (E.g., the embeddings dropdown only lists models the Ollama instance reports as embedding-capable.)",
        "Migration safety: when a user switches embedding model, the system re-embeds the corpus in the background with a progress indicator. Old vectors are kept namespaced until the new index is complete.",
    ],
    [
        "Android Settings → AI screens (Compose).",
        "Web Settings → AI pages.",
        "Edge function `ai/test-connection` (V1: calls Ollama /api/tags; V2: provider-aware).",
        "Edge function `ai/usage` for the usage dashboard.",
        "DB migration adding `ai_keys` (encrypted, unused in V1), `ai_usage_log`, expanding `users_profile.ai_prefs`.",
        "Re-embed background job triggered on embedding-model change.",
    ],
    "M1, M2, M9.",
    "Opus for the routing-config schema (must survive V2 multi-provider without redesign); Sonnet for the Settings UI screens; Haiku for the test-connection edge function.")

module_block(doc, "M14", "DevOps, Observability & Release",
    "GitHub Actions pipelines, environment management, Sentry, PostHog, and Play "
    "Store release wiring. The 'unglamorous but mandatory' module.",
    """// .github/workflows/
//   ci.yml         (lint + test on every push)
//   deploy-web.yml (Vercel deploy on main)
//   build-apk.yml  (signed AAB upload to Play internal track)
//   db-migrate.yml (supabase db push on tag)""",
    [
        "Three environments: dev (local), staging (separate Supabase project), prod.",
        "Secrets in GitHub Encrypted Secrets and Vercel/Supabase env config (never in repo).",
        "Sentry source maps uploaded on every web deploy; ProGuard mapping uploaded on every Android build.",
        "PostHog: anonymous product analytics; opt-out toggle in settings.",
    ],
    [
        "All four pipelines green on a sample PR.",
        "Runbook: 'How to roll back', 'How to rotate Claude API key', 'How to add a new user'.",
    ],
    "M1–M13.",
    "Haiku — config-heavy, mechanical, well-documented externally.")

page_break(doc)


# ==================== 8. DATA MODEL ====================
add_heading(doc, "8. Data Model", level=1)
add_para(doc,
    "Postgres schema (Supabase). All tables include `id uuid pk default gen_random_uuid()`, "
    "`owner_id uuid not null references auth.users(id)`, `created_at`, `updated_at`. "
    "Row-Level Security enforces `owner_id = auth.uid()` on every table from day one — "
    "this is what makes multi-tenant trivial later.")

add_table(doc,
    headers=["Table", "Key Columns", "Notes"],
    rows=[
        ["users_profile", "user_id pk → auth.users, storage_provider, calendar_provider, ai_prefs jsonb", "1 row per user; holds per-user provider choices."],
        ["notebooks", "id, owner_id, name, color, sort_order", "One level deep."],
        ["tags", "id, owner_id, name, color", "Unique on (owner_id, lower(name))."],
        ["notes", "id, owner_id, notebook_id, title, body_json jsonb, body_tsv tsvector, is_pinned, is_trashed, trashed_at", "tsvector maintained by trigger for keyword search."],
        ["note_tags", "note_id, tag_id (composite pk)", "Join table."],
        ["attachments_refs", "id, owner_id, note_id, provider, external_id, display_name, mime, size_bytes", "References user's Drive/OneDrive; never bytes."],
        ["events_mirror", "id, owner_id, external_event_id, calendar_id, title, start_at, end_at, payload jsonb", "Read-only mirror of Google Calendar."],
        ["note_event_links", "note_id, external_event_id", "Tracks events created from a note."],
        ["note_embeddings", "id, owner_id, note_id, chunk_index, content text, embedding vector, embedding_namespace text", "HNSW index per namespace; namespace = provider:model so multiple embedding models can coexist during migration."],
        ["ai_conversations", "id, owner_id, title, model_hint, created_at, last_message_at, archived_summary text", "Persistent chat threads; archived_summary holds the compressed older turns."],
        ["ai_messages", "id, conversation_id, role, content, citations jsonb, provider, model, prompt_tokens, completion_tokens, cost_usd, created_at", "One row per turn; includes which provider/model served it."],
        ["ai_memory", "id, owner_id, fact text, source ('user' | 'derived'), confidence, created_at, last_used_at", "Long-term facts about the user; prepended to new conversations."],
        ["ai_keys", "owner_id, provider, encrypted_secret bytea, endpoint_url text", "Encrypted in Supabase Vault; decrypted only at call time."],
        ["ai_usage_log", "id, owner_id, provider, model, task, prompt_tokens, completion_tokens, cost_usd, latency_ms, at", "Powers the Settings → AI → Usage dashboard."],
        ["reminders", "id, owner_id, note_id, fire_at, payload jsonb, status", "Local + (future) push."],
        ["sync_outbox (Android local only)", "id, op, payload jsonb, attempts, created_at", "Room table; never synced to backend."],
        ["audit_log", "id, owner_id, actor, action, target, at, meta jsonb", "Append-only; powers 'recently changed' and forensic review."],
    ],
    col_widths=[1.6, 2.8, 2.6])

add_callout(doc, "RLS policy template",
    "create policy \"owner_can_all\" on <table> for all using (owner_id = auth.uid()) "
    "with check (owner_id = auth.uid()); — applied uniformly. Public launch needs zero "
    "schema changes; only feature flags.")

page_break(doc)


# ==================== 9. API ====================
add_heading(doc, "9. API Contract (Backend ↔ Clients)", level=1)

add_para(doc,
    "Two API surfaces: (a) auto-generated REST via Supabase PostgREST for CRUD, and "
    "(b) custom Edge Functions for orchestrated operations (AI, calendar sync, OAuth "
    "token refresh). Clients always authenticate with a Supabase JWT.")

add_heading(doc, "9.1 PostgREST (auto)", level=2)
add_code_block(doc, """GET    /rest/v1/notes?select=*&order=updated_at.desc&limit=50
POST   /rest/v1/notes                    body: { title, body_json, notebook_id }
PATCH  /rest/v1/notes?id=eq.<uuid>       body: { title?, body_json?, is_pinned?, is_trashed? }
DELETE /rest/v1/notes?id=eq.<uuid>       (rarely used — prefer is_trashed = true)

GET    /rest/v1/notebooks?select=*
GET    /rest/v1/tags?select=*
GET    /rest/v1/events_mirror?start_at=gte.<iso>&start_at=lte.<iso>
""")

add_heading(doc, "9.2 Custom Edge Functions", level=2)
add_table(doc,
    headers=["Endpoint", "Method", "Purpose"],
    rows=[
        ["/functions/v1/ai/chat", "POST (SSE)", "Stream chat-with-notes answer."],
        ["/functions/v1/ai/summarize", "POST", "Return summary of a note."],
        ["/functions/v1/ai/suggest-tags", "POST", "Return tag suggestions."],
        ["/functions/v1/ai/suggest-title", "POST", "Return title suggestions."],
        ["/functions/v1/embeddings/index", "POST (internal)", "Re-embed a note. Called by DB trigger."],
        ["/functions/v1/calendar/sync", "POST", "Refresh events_mirror for a date range."],
        ["/functions/v1/calendar/create-event", "POST", "Create a Calendar event from a note."],
        ["/functions/v1/storage/sign", "POST", "Return a short-lived signed URL for an attachment."],
        ["/functions/v1/auth/refresh-provider-token", "POST", "Refresh Google/Microsoft tokens server-side."],
        ["/functions/v1/export/full", "POST", "Build and return a .zip of the user's full corpus."],
    ],
    col_widths=[2.7, 1.0, 2.8])

add_heading(doc, "9.3 Realtime channels", level=2)
add_code_block(doc, """// Supabase Realtime
//   channel: per-user, postgres_changes on tables
//     notes, notebooks, tags, note_tags, attachments_refs
// Clients subscribe on auth and unsubscribe on sign-out.
""")

page_break(doc)


# ==================== 10. AI DESIGN ====================
add_heading(doc, "10. Second-Brain AI Design (Multi-Provider)", level=1)

add_para(doc,
    "This section is the contract for what makes the system a second brain rather than "
    "a notes app with AI bolted on. It also locks down how multi-provider works: any "
    "supported provider can serve any task, with the user in control of routing.")

add_heading(doc, "10.1 Model capability matrix (Ollama models on RTX 3060 12GB)", level=2)
add_para(doc,
    "V1 ships with Ollama only. The matrix below documents which models on the user's "
    "Legion are recommended for which tasks. The UI in M15 reads model metadata from the "
    "Ollama endpoint and populates dropdowns accordingly. When V2 ships additional "
    "providers, this matrix gains columns; the schema is already provider-aware.")

add_table(doc,
    headers=["Model (Ollama)", "Size", "VRAM use", "Best for", "Speed on RTX 3060"],
    rows=[
        ["llama3.1:8b", "8B params, 4.7GB Q4", "~7GB", "Chat, Q&A, summarization, briefing — DEFAULT for chat task", "30–60 tok/s"],
        ["qwen2.5:7b", "7B, 4.4GB Q4", "~6GB", "Alternative to llama3.1:8b; better at structured output", "35–65 tok/s"],
        ["llama3.2:3b", "3B, 2GB Q4", "~3GB", "Tagging, titling, action-item extraction — DEFAULT for bulk tasks", "70–120 tok/s"],
        ["phi3.5:3.8b", "3.8B, 2.2GB Q4", "~3GB", "Alternative bulk-task model; strong reasoning for its size", "60–100 tok/s"],
        ["nomic-embed-text", "137M, 274MB", "~0.5GB", "Embeddings (768-dim) — DEFAULT for embed task", "Batch ~1000 texts/sec"],
        ["mxbai-embed-large", "335M, 670MB", "~1GB", "Higher-quality embeddings (1024-dim); slower", "Batch ~500 texts/sec"],
        ["llava:7b", "7B, 4.7GB Q4", "~7GB", "Vision (V2 — when image-attachment understanding is added)", "20–40 tok/s"],
        ["whisper.cpp (separate)", "small=466MB / medium=1.5GB", "CPU or 1–2GB GPU", "Voice transcription (V1.5)", "Real-time or faster"],
    ],
    col_widths=[1.5, 0.9, 0.8, 2.3, 1.1])

add_callout(doc, "Hardware fit",
    "The Legion (RTX 3060 12GB) can run llama3.1:8b + nomic-embed-text concurrently with "
    "VRAM to spare. Ollama keeps models loaded for ~5 min after last use; we'll tune "
    "OLLAMA_KEEP_ALIVE=30m so chat feels instant.")

add_heading(doc, "10.2 RAG pipeline for chat-with-notes", level=2)
add_code_block(doc, """User question
   │
   ▼
Resolve routing for task=CHAT       → e.g., (claude, claude-sonnet-4-6)
Resolve routing for task=EMBED      → e.g., (voyage, voyage-3-large)
   │
   ▼
Embed question via the embedding adapter
   │
   ▼
Vector search (pgvector top-k=8, filter by user_id, optional notebook/tag,
              embedding_namespace = active embedding model)
   │
   ▼
Load prior conversation turns (last N, summarized if long) for memory
   │
   ▼
Build prompt (provider-neutral structure; adapter formats per API):
   [System]   You are the user's second brain. Answer using the provided
              notes and conversation history. Always cite [[NoteId:xxx]].
   [Cached]   Few-shot examples (5 turns) — cached on Claude/OpenAI
   [Memory]   Compressed prior turns
   [Context]  Top-k chunks with NoteId tags
   [User]     The actual question
   │
   ▼
Active chat adapter, streaming
   │
   ▼
Client renders tokens; on completion, parses [[NoteId:xxx]] citations
into clickable links; conversation persisted to ai_messages.""")

add_heading(doc, "10.3 Persistent memory & conversations", level=2)
add_para(doc,
    "Second-brain implies continuity — the AI should remember what you talked about "
    "yesterday. Implementation:")
for item in [
    "Conversations are stored in `ai_conversations` (id, title, created_at, last_message_at) and `ai_messages` (conversation_id, role, content, citations, tokens, cost).",
    "When a conversation grows beyond a provider's context window, the oldest turns are summarized by Haiku (or the cheapest configured model) and the raw turns are archived. The active conversation always fits in context.",
    "Cross-conversation memory: the system extracts long-term facts ('user is working on Project Phoenix', 'user lives in Pune', 'user prefers concise bullet answers') into a per-user `ai_memory` table. These facts are prepended to every new conversation.",
    "User can view, edit, or delete any memory item from Settings → AI → Memory.",
]:
    add_bullet(doc, item)

add_heading(doc, "10.4 Prompt portability across providers", level=2)
add_para(doc,
    "The same logical prompt must work on Claude, OpenAI, and Ollama. Strategy:")
for item in [
    "Prompts are authored in a provider-neutral template format with placeholders.",
    "Each adapter has a `format(prompt: NeutralPrompt): ProviderRequest` method that handles the API-specific wrapping (Claude's system/messages structure, OpenAI's chat-completions structure, Ollama's prompt template).",
    "Provider-specific tuning (e.g., explicit 'respond in JSON' for Ollama models that don't reliably do JSON mode) lives in the adapter as a per-model override, not in the shared prompt.",
    "All prompts live in `/prompts/*.md` and are versioned; A/B testing prompts across providers is supported via M15's routing config.",
]:
    add_bullet(doc, item)

add_heading(doc, "10.5 Ollama integration specifics (V1)", level=2)
for item in [
    "User installs Ollama on the Legion (Windows native installer). Ollama runs as a background service on port 11434.",
    "User installs Tailscale on the Legion and on the Android phone (and laptop). The Legion gets a stable Tailscale hostname like `legion.<tailnet>.ts.net`.",
    "Settings → AI → Endpoint defaults to `http://legion.<tailnet>.ts.net:11434` — the user just confirms.",
    "Backend Edge Functions are also added to the user's tailnet via Tailscale's GitHub Actions / Vercel integration (or via a serverless-friendly auth header pattern documented in §17).",
    "Pure-LAN mode (toggle): when the phone is on home Wi-Fi, calls go direct LAN bypassing the backend entirely — lower latency and full privacy.",
    "Model discovery: on 'Test connection', the adapter calls /api/tags to enumerate installed models and populates routing dropdowns.",
    "If Ollama is unreachable (Legion off, Tailscale disconnected): the system fails loudly with an actionable error ('Cannot reach your Ollama at <host>. Is the desktop on?'). V2 will offer optional fallback to a cloud provider.",
    "Keep-alive tuning: backend issues a /api/generate with `keep_alive: '30m'` so the model stays warm; cold-start of llama3.1:8b is ~5s, warm is sub-100ms to first token.",
]:
    add_bullet(doc, item)

add_heading(doc, "10.6 Daily briefing, related notes, voice capture", level=2)
add_para(doc, "Daily briefing", bold=True)
for item in [
    "Triggered by pg_cron at the user's configured time; runs as a backend job (no client needed).",
    "Aggregates: yesterday's new/edited notes, today's calendar events, open action items, anniversaries of older notes ('on this day last year').",
    "AI synthesizes a 5-bullet Markdown briefing using the routed model for 'briefing' task.",
    "Delivered as a system note (pinned to top of Today view) and an Android push notification.",
]:
    add_bullet(doc, item)
add_para(doc, "Related notes (connection discovery)", bold=True)
for item in [
    "Pure vector-similarity — no LLM call — so latency is sub-100ms.",
    "Shown as a side panel in the editor and a tappable section on Android.",
    "Updates live as the user types (debounced 2s); helps surface 'I wrote about this before' moments.",
]:
    add_bullet(doc, item)
add_para(doc, "Voice capture (V1.5 — deferred from V1)", bold=True)
for item in [
    "Persistent quick-tile in Android notification shade for one-tap recording.",
    "Audio uploaded to a small whisper.cpp service running on the Legion (alongside Ollama).",
    "Transcript fed back to the chat model (llama3.1:8b) which suggests notebook, tags, and title.",
    "User reviews on a one-screen card and taps Save.",
    "End-to-end target: under 8s from stop-recording to confirmed note for a 30-second clip on the Legion.",
    "Skipped in V1 to keep scope tight; will be added in P4.5 once the rest is dogfooded.",
]:
    add_bullet(doc, item)

add_heading(doc, "10.7 Performance guardrails (V1)", level=2)
for item in [
    "Cost is effectively $0 in V1 (local inference). Budget cap UI is built but inert until V2 cloud providers are added.",
    "Per-task model defaults favor the smallest model that does the task acceptably (3B for tagging, 8B for chat).",
    "Embedding only on note save, debounced 10s after last edit; not on every keystroke.",
    "Top-k retrieval truncated to 4–8 chunks; never dump the whole corpus into context.",
    "Telemetry visible in real time (Settings → AI → Usage): latency, tokens, model, task.",
    "If chat latency exceeds 3s to first token, the UI shows 'Warming model…' so the user knows it's a cold start, not a hang.",
]:
    add_bullet(doc, item)

add_heading(doc, "10.8 Privacy & user controls", level=2)
for item in [
    "V1 architectural guarantee: note content never traverses the public internet for AI processing. The path is: backend (Supabase Edge Function) → Tailscale tunnel → Legion → Ollama. Tailscale is end-to-end encrypted; Ollama runs entirely on the user's hardware.",
    "AI calls only fire on explicit user action OR background indexing (embeddings) of newly-saved notes.",
    "Settings → AI → 'Disable AI features' toggle. When off, no AI calls happen at all.",
    "Notes can be marked 'AI-excluded'; they are skipped during embedding and retrieval.",
    "Per-notebook AI policy: e.g., 'Personal Journal' can be flagged 'local-only' (today this is trivially satisfied since everything is local; in V2 when cloud providers exist, this flag will prevent them from touching that notebook).",
    "Full audit log per AI call (model, tokens, latency, task) viewable by the user.",
]:
    add_bullet(doc, item)

page_break(doc)


# ==================== 11. SECURITY ====================
add_heading(doc, "11. Security, Privacy & Compliance", level=1)

add_table(doc,
    headers=["Area", "Control", "Notes"],
    rows=[
        ["Transport", "TLS 1.2+ for all client↔backend traffic.", "Enforced by Supabase and Vercel."],
        ["Data at rest", "Supabase Postgres encryption at rest (AES-256).", "Default."],
        ["Authn", "Google OAuth via Supabase Auth.", "No password storage anywhere."],
        ["Authz", "Postgres Row-Level Security on every table.", "Tested with policy unit tests."],
        ["Provider tokens", "Stored in Supabase Vault, never in client storage beyond the access-token cookie.", "Refresh tokens server-side only."],
        ["Attachments", "Stored in user's own Drive/OneDrive.", "Platform never sees binary bytes."],
        ["Secrets", "Anthropic key, Google client secret, Microsoft client secret — Supabase Vault + GitHub Encrypted Secrets.", "Rotation runbook in M14."],
        ["Logging", "Server logs scrub note bodies and personal info.", "PII budget: zero."],
        ["Backups", "Supabase daily backups (managed); manual full-corpus export available on demand.", "Export endpoint in §9.2."],
        ["Deletion", "Trash retention 30 days; user-initiated full-account delete cascades all rows + revokes provider tokens.", "GDPR-aligned."],
        ["E2EE", "Out of scope V1. Planned for V3.", "Schema designed to allow note body to become opaque ciphertext later."],
    ],
    col_widths=[1.4, 3.0, 2.6])

page_break(doc)


# ==================== 12. ROADMAP ====================
add_heading(doc, "12. Phased Roadmap", level=1)

add_table(doc,
    headers=["Phase", "Scope", "Modules", "Duration"],
    rows=[
        ["P0 — Foundation", "Repos, CI/CD, Supabase project, domain model, auth, basic notes CRUD on web and Android.",
         "M1, M2, M4, M14", "2–3 weeks"],
        ["P1 — Editor & sync", "Tiptap web editor, Compose Android editor, offline sync, notebooks, tags, keyword search.",
         "M5, M6", "3–4 weeks"],
        ["P2 — Attachments", "Google Drive adapter + attachment pipeline.",
         "M3, M7", "1–2 weeks"],
        ["P3 — Second-brain AI core (Ollama only)", "Ollama set up on the Legion (§17). Tailscale tunnel. Embeddings + RAG + chat with persistent conversations + summarize/tag/title. AI Settings UI (endpoint, model picker, routing).",
         "M9 (Ollama adapter), M10, M15", "3 weeks"],
        ["P4 — Second-brain advanced", "Related-notes panel, daily briefing, AI memory, cost/usage dashboard (with $0 in V1).",
         "M9 (extensions)", "2 weeks"],
        ["P4.5 — Voice capture (optional)", "whisper.cpp on the Legion; Android quick-tile; transcript → classified note flow.",
         "M9 (voice)", "1 week"],
        ["P5 — Calendar & reminders", "Google Calendar two-way; local reminders.",
         "M8, M11", "1–2 weeks"],
        ["P6 — Polish & dogfood", "Personal cutover from Evernote; bug fixing; observability.",
         "All", "2 weeks"],
        ["V2 (later) — Multi-provider", "Add ClaudeAdapter and OpenAIAdapter (one file each). Expose Provider column in Settings → AI → Routing. Reuse existing M15 schema unchanged.",
         "(adapters only)", "1 week"],
        ["V3 (later) — Public launch prep", "OneDrive adapter, Outlook adapter, billing/plans, E2EE, iOS app.",
         "(new modules)", "Open"],
    ],
    col_widths=[1.5, 2.7, 1.4, 1.0])

add_callout(doc, "Total to personal cutover",
    "≈ 12–15 weeks at part-time pace using AI developers per the assignment in Section 13. "
    "Ollama-only V1 is 3 weeks faster than the original multi-provider V1 because only one "
    "AI adapter ships. The architecture cost of supporting future providers is paid up front "
    "in the port design (M9 / M15), not in shipping multiple adapters now.")

page_break(doc)


# ==================== 13. WORK PACKAGES ====================
add_heading(doc, "13. Work-Package Distribution (Opus / Sonnet / Haiku)", level=1)

add_para(doc,
    "Each module from Section 7 is sized and matched to a model based on the kind of "
    "thinking it requires. The rule of thumb: Opus for design judgment, Sonnet for "
    "well-specified implementation, Haiku for mechanical / config / boilerplate work.")

add_table(doc,
    headers=["Module", "Model", "Why this model", "Hand-off Artefacts"],
    rows=[
        ["M1 Domain Model", "Sonnet", "Mechanical translation of spec → types.", "§7.M1 + ER diagram in §8"],
        ["M2 Auth", "Opus (design) + Sonnet (adapters)", "Provider token strategy is judgment; adapter code isn't.", "§7.M2 + §11"],
        ["M3 Storage Provider", "Opus", "OAuth scopes, quota behavior, and dual-adapter contract need careful design.", "§7.M3 + provider docs links in §16"],
        ["M4 Notes Core", "Sonnet", "Schema + RLS + CRUD endpoints are well-specified.", "§7.M4 + §8 + §9"],
        ["M5 Sync Engine", "Opus", "Concurrency, idempotency, conflict policy — highest-judgment module.", "§7.M5 + 'How sync works' explainer"],
        ["M6 Editor (web)", "Sonnet", "Tiptap is well-trodden.", "§7.M6 + frozen schema doc"],
        ["M6 Editor (Android)", "Opus", "Compose rich-text editing is non-trivial.", "§7.M6 + frozen schema doc"],
        ["M7 Attachment Pipeline", "Sonnet", "Bounded; thumbnail logic is mechanical.", "§7.M7"],
        ["M8 Calendar Integration", "Sonnet", "Google Calendar API is well-documented; adapter pattern already designed.", "§7.M8"],
        ["M9 Second-Brain AI core", "Opus", "RAG quality, prompt design for open-weight models (which need more careful prompting), conversation compression, memory extraction — highest-judgment work.", "§7.M9 + §10 + prompt library"],
        ["M9 — Ollama adapter (V1)", "Sonnet", "Mechanical once port is designed; handles endpoint reachability and model discovery.", "§7.M9 + §10.5 + §17"],
        ["M9 — Claude adapter (V2)", "Sonnet", "Deferred. One file when V2 happens.", "§7.M9 (port) + Anthropic API docs"],
        ["M9 — OpenAI adapter (V2)", "Sonnet", "Deferred. One file when V2 happens.", "§7.M9 (port) + OpenAI API docs"],
        ["M10 Embedding/Vector", "Sonnet", "Well-bounded; namespace strategy documented.", "§7.M10"],
        ["M11 Notifications", "Haiku", "Small, mechanical, platform-API driven.", "§7.M11"],
        ["M12 Android Shell", "Opus (scaffold) + Sonnet (screens)", "DI graph + navigation is design; screens are implementation.", "§7.M12"],
        ["M13 Web Shell", "Sonnet", "Modern Next.js is well-trodden.", "§7.M13"],
        ["M14 DevOps & Release", "Haiku", "Config-heavy, mechanical, well-documented externally.", "§7.M14"],
        ["M15 AI Settings & Routing", "Opus (design) + Sonnet (UI) + Haiku (test-connection funcs)", "Routing config schema is consequential; UI is well-trodden; test-connection is trivial.", "§7.M15 + §10"],
    ],
    col_widths=[1.4, 1.5, 2.6, 1.5])

add_heading(doc, "13.1 How to hand a module to an AI developer", level=2)
add_para(doc,
    "Each hand-off prompt should contain: (a) the relevant Section 7 sub-block "
    "verbatim, (b) the relevant Data Model rows (Section 8), (c) any API contracts "
    "the module touches (Section 9), (d) the explicit 'Definition of Done' below. "
    "Nothing else. The architecture document is the single source of truth.")

add_heading(doc, "13.2 Definition of Done (per module)", level=2)
for item in [
    "All public functions on the port interface are implemented.",
    "Unit tests cover happy path + one failure path per public function.",
    "An integration test exercises the module end-to-end against a real adapter (or a documented fake).",
    "No code outside the module's adapter files references the underlying provider by name.",
    "README in the module folder explaining responsibilities, ports, and how to swap the adapter.",
]:
    add_bullet(doc, item)

page_break(doc)


# ==================== 14. FORWARD COMPAT ====================
add_heading(doc, "14. Forward Compatibility & Multi-Tenant Path", level=1)

add_para(doc, "Five design decisions make the personal-to-SaaS transition cheap:")

add_numbered(doc,
    "Every table carries owner_id from day one. RLS is enforced even with a single "
    "user. Adding more users requires only signup flow + billing — no migration.")
add_numbered(doc,
    "Every external dependency is a port. To support new providers (OneDrive, "
    "Outlook, OpenAI as a fallback) we add an adapter, not a fork.")
add_numbered(doc,
    "Per-user provider choice is stored in users_profile. Different users can use "
    "different storage and calendar providers without any plumbing change.")
add_numbered(doc,
    "All AI prompts and embeddings include explicit user_id filters even though "
    "RLS would also enforce them. Defense in depth against a future bug.")
add_numbered(doc,
    "Edge Functions are stateless. Moving from Supabase Edge Functions to a "
    "standalone Node service (Fly.io, Render, AWS) is a re-deploy, not a rewrite.")

add_heading(doc, "14.1 What will need to change for public launch", level=2)
for item in [
    "Billing — Stripe + plan gating in users_profile.",
    "Quotas — per-user note count, attachment storage, AI tokens-per-day.",
    "Admin dashboard — minimal Next.js admin app over the same Supabase project.",
    "Support — error report visibility, account recovery, email transactional (Resend/Postmark).",
    "iOS app — separate codebase reusing the same backend and prompt library.",
    "E2EE — note body becomes encrypted ciphertext; embeddings move client-side (deferred V3).",
    "Compliance — DPA, privacy policy, ToS, SOC 2 path if needed.",
]:
    add_bullet(doc, item)

page_break(doc)


# ==================== 15. RISKS ====================
add_heading(doc, "15. Risks, Tradeoffs & Open Decisions", level=1)

add_table(doc,
    headers=["#", "Item", "Risk / Tradeoff", "Mitigation / Resolution"],
    rows=[
        ["R1", "Two clients, two editors", "Web + Android each need a rich-text editor; Compose rich-text is harder than Tiptap.",
         "Freeze the document schema first (M6). Android can start with edit-in-Markdown + render-rich and upgrade incrementally."],
        ["R2", "Drive/OneDrive quota", "Personal Google accounts share 15 GB across Drive/Gmail/Photos.",
         "Show a quota gauge in settings; default attachment size cap configurable; thumbnail-only originals option."],
        ["R3", "pgvector at scale", "HNSW is fine for single-user, but multi-tenant with 1M+ vectors may need partitioning.",
         "Port abstracts this. Pinecone/Qdrant adapter is a V3 task, not V1."],
        ["R4", "Sync conflicts", "Last-writer-wins loses data if two devices edit offline.",
         "Stage conflict body in a 'Conflicts' notebook for user review (not silent drop). Documented in M5."],
        ["R5", "Anthropic API cost overrun", "Curiosity-driven chat could blow the personal budget.",
         "Per-day cap (§10.3) plus model routing. Telemetry in PostHog to see actual cost."],
        ["R6", "Vendor lock-in (Supabase)", "Migrating off managed Supabase later is non-trivial.",
         "Standard Postgres + standard SQL + Edge Functions in plain TS. Self-host path is documented; not exercised V1."],
        ["R7", "iOS users post-launch", "Building iOS is a second codebase.",
         "Acknowledged tradeoff. Kotlin Multiplatform considered for V3 to share business logic."],
        ["R8", "Ollama outage (Legion off, restarted, Tailscale dropped)", "Without fallback, the second brain is unavailable when the Legion is offline.",
         "Settings shows endpoint health prominently. Clear error messaging. V2 multi-provider unlocks optional cloud fallback (architecture-ready)."],
        ["R9", "Open-weight model quality below Claude/GPT-4o", "Q&A and synthesis may be noticeably weaker than cloud frontier models.",
         "Accept the tradeoff for V1 (cost, privacy). Prompt library tuned for open-weight models. If quality is unacceptable after dogfood, V2 ClaudeAdapter is a one-file fix."],
        ["R10", "Embedding model lock-in", "Switching embedding model later requires re-embedding the whole corpus.",
         "Embeddings namespaced by model in DB (§8). M15's re-embed background job is built in V1. Painful but solved."],
        ["R11", "Tailscale dependency", "Tailscale free tier has device limits (currently 100); if Tailscale changes policy, reachability strategy needs to change.",
         "Cloudflare Tunnel adapter documented as alternative (§17.5). Endpoint URL is configurable so swap is a Settings change."],
        ["R12", "RTX 3060 model size ceiling", "Cannot run 70B-class models that exceed Sonnet/GPT-4o quality.",
         "Accept the ceiling for V1. The architecture supports 'route some tasks to cloud, others to Ollama' in V2 for users who want both."],
        ["R13", "Legion power/electricity", "Always-on Legion has non-trivial idle draw.",
         "Use Windows power management to spin GPU down when idle; Ollama loads models on demand. Idle draw with GPU asleep is ~30W; under load ~250W."],
    ],
    col_widths=[0.4, 1.7, 2.5, 2.4])

add_heading(doc, "15.1 Decisions explicitly deferred (require user input)", level=2)
add_bullet(doc, "Domain/branding for public launch — not needed for personal V1.")
add_bullet(doc, "iOS timeline — depends on outcome of personal V1.")
add_bullet(doc, "E2EE design — depends on multi-device threat model at launch time.")
add_bullet(doc, "Billing model (subscription, lifetime, freemium) — to be decided at P6.")

page_break(doc)


# ==================== 16. APPENDIX ====================
add_heading(doc, "16. Appendix: Glossary & References", level=1)

add_heading(doc, "16.1 Glossary", level=2)
add_table(doc,
    headers=["Term", "Meaning"],
    rows=[
        ["Port", "An interface defining a capability the application core needs. Provider-agnostic."],
        ["Adapter", "A concrete implementation of a port against a specific provider (e.g., GoogleDriveAdapter)."],
        ["RAG", "Retrieval-Augmented Generation. Retrieves relevant context from the user's data, then asks the LLM to answer using only that context."],
        ["RLS", "Row-Level Security. Postgres mechanism that enforces per-row access in the database itself."],
        ["pgvector", "Postgres extension for storing and searching embedding vectors."],
        ["HNSW", "Hierarchical Navigable Small World. The vector index algorithm pgvector uses for fast approximate nearest-neighbor search."],
        ["DXA", "Document Word XML measurement unit (1440 DXA = 1 inch). Used internally by Word."],
        ["KMP", "Kotlin Multiplatform. Future option to share business logic between Android and iOS."],
        ["SSE", "Server-Sent Events. Used to stream Claude responses from edge functions to clients."],
    ],
    col_widths=[1.5, 5.0])

add_heading(doc, "16.2 References (for the dev team)", level=2)
for item in [
    "Anthropic API documentation (Messages API, prompt caching, tool use).",
    "Supabase docs: Auth, Realtime, Edge Functions, pgvector guide.",
    "Google Drive API v3 reference and OAuth 2.0 scopes guide.",
    "Microsoft Graph API: OneDrive endpoints + MSAL for Kotlin/JS.",
    "Google Calendar API v3 + push-notification channels.",
    "Tiptap docs + ProseMirror schema guide.",
    "Jetpack Compose docs, particularly Navigation and state management.",
    "Voyage AI embeddings documentation (voyage-3-large).",
]:
    add_bullet(doc, item)

add_heading(doc, "16.3 Document control", level=2)
add_table(doc,
    headers=["Version", "Date", "Author", "Change"],
    rows=[
        ["1.0", "2026-05-19", "V. Gakas (with Claude)", "Initial draft for review."],
    ],
    col_widths=[0.8, 1.3, 2.4, 2.5])


page_break(doc)


# ==================== 17. OLLAMA DEPLOYMENT GUIDE ====================
add_heading(doc, "17. Ollama Deployment Guide (V1 Hardware-Specific)", level=1)

add_para(doc,
    "This section is a concrete deployment guide for the specific V1 hardware target: "
    "a Lenovo Legion desktop running Windows, Intel i7, NVIDIA RTX 3060 (12GB VRAM), "
    "16GB system RAM. It is intentionally prescriptive rather than abstract — copy the "
    "commands and the system works.")

add_heading(doc, "17.1 Step 1 — Install Ollama on the Legion", level=2)
add_code_block(doc, """# Windows installer
# Download from: https://ollama.com/download/windows
# Run OllamaSetup.exe — installs Ollama as a background service on port 11434

# Verify after install (PowerShell):
ollama --version
curl http://localhost:11434/api/tags     # should return JSON

# Allow remote access (so backend can reach via Tailscale):
# Set environment variable OLLAMA_HOST=0.0.0.0:11434
# Settings → System → About → Advanced system settings → Environment Variables
# Then restart the Ollama service (or reboot).""")

add_heading(doc, "17.2 Step 2 — Pull the V1 model set", level=2)
add_code_block(doc, """# Chat / Q&A (default routing for task=CHAT)
ollama pull llama3.1:8b           # ~4.7 GB, fits in 12GB VRAM with headroom

# Bulk tasks (default routing for task=TAG, TITLE, ACTIONS)
ollama pull llama3.2:3b           # ~2 GB, runs at 70–120 tok/s

# Embeddings (default routing for task=EMBED)
ollama pull nomic-embed-text      # ~270 MB, 768-dim vectors

# Verify
ollama list""")

add_heading(doc, "17.3 Step 3 — Tune Ollama for low-latency chat", level=2)
add_para(doc,
    "Set these environment variables before starting Ollama (System → Environment "
    "Variables on Windows):")
add_table(doc,
    headers=["Variable", "Value", "Why"],
    rows=[
        ["OLLAMA_HOST", "0.0.0.0:11434", "Listen on all interfaces so Tailscale can reach it."],
        ["OLLAMA_KEEP_ALIVE", "30m", "Keep models loaded in VRAM for 30 minutes after last use — chat feels instant."],
        ["OLLAMA_NUM_PARALLEL", "2", "Allow concurrent requests (e.g., chat + background embedding)."],
        ["OLLAMA_MAX_LOADED_MODELS", "3", "Keep up to 3 models warm (chat-8b, bulk-3b, embed); fits in 12GB."],
        ["OLLAMA_FLASH_ATTENTION", "1", "Faster inference on RTX 3060 (CUDA-capable)."],
    ],
    col_widths=[2.0, 1.5, 3.0])

add_heading(doc, "17.4 Step 4 — Install Tailscale", level=2)
add_code_block(doc, """# On the Legion:
#   Install from https://tailscale.com/download/windows
#   Sign in (free tier — up to 100 devices, more than enough)
#   Note the Legion's Tailscale name, e.g.:  legion.<your-tailnet>.ts.net

# On the Android phone:
#   Install Tailscale from Play Store
#   Sign in with the same account
#   Connect (toggle on)

# Verify from phone:
#   In any browser on the phone: http://legion.<tailnet>.ts.net:11434
#   Should return: "Ollama is running"

# On any other device (laptop, etc.) where you'll use the web app:
#   Install Tailscale, sign in, connect""")

add_heading(doc, "17.5 Step 5 — Connect the backend (Supabase Edge Functions)", level=2)
add_para(doc,
    "Supabase Edge Functions run on Deno Deploy. To reach the Legion via Tailscale, we "
    "have three options — choose one:")
add_table(doc,
    headers=["Option", "How", "Tradeoff"],
    rows=[
        ["A. Tailscale Funnel (recommended)", "Run `tailscale funnel 11434` on the Legion. This exposes Ollama at a TLS-terminated public URL like https://legion.<tailnet>.ts.net (HTTPS only, your tailnet only).", "Simplest. Free. No extra services."],
        ["B. Cloudflare Tunnel", "Install cloudflared on the Legion; create a tunnel pointing to localhost:11434.", "Slightly more setup; provides a stable HTTPS URL independent of Tailscale."],
        ["C. Backend-on-tailnet", "Replace Supabase Edge Functions with a small Node service hosted on a Tailscale-connected VPS.", "More flexibility; more moving parts. Only worth doing if you outgrow Edge Functions."],
    ],
    col_widths=[1.8, 3.2, 1.5])

add_callout(doc, "Recommended for V1",
    "Option A (Tailscale Funnel). Two commands on the Legion, no other infrastructure, "
    "free, and the Ollama endpoint URL the backend uses becomes simply "
    "https://legion.<your-tailnet>.ts.net — plug that into Settings → AI → Endpoint and "
    "you're done.")

add_heading(doc, "17.6 Step 6 — Whisper for voice capture (V1.5, optional)", level=2)
add_code_block(doc, """# Option A: whisper.cpp as a separate process on the Legion
#   git clone https://github.com/ggerganov/whisper.cpp
#   build with CUDA flags for GPU acceleration
#   download model: bash ./models/download-ggml-model.sh medium
#   run server: ./server -m models/ggml-medium.bin --port 11435

# Option B: speaches (whisper-compatible OpenAI API wrapper, Docker)
#   docker run -p 11435:8000 fedirz/faster-whisper-server:latest-cuda

# Settings → AI → Voice → endpoint = http://legion.<tailnet>.ts.net:11435
# The Voice adapter calls /audio/transcriptions (OpenAI-compatible).""")

add_heading(doc, "17.7 Operational notes", level=2)
for item in [
    "Windows updates that require restart: the Ollama service will resume automatically; chat will have a 5–10s cold start on first query after reboot.",
    "If the Legion's IP on the tailnet changes (rare), the hostname stays stable — no Settings change needed.",
    "Monitoring: use Task Manager → Performance → GPU on the Legion to confirm Ollama is using CUDA. If you see 100% CPU and 0% GPU, the model is running CPU-only — check NVIDIA drivers.",
    "VRAM check: with llama3.1:8b loaded, expect 7–8 GB used; with all three models warm, expect 10–11 GB used. Restart Ollama if you see VRAM exhaustion errors.",
    "Backup: Ollama models live in `%USERPROFILE%\\.ollama\\models`. Not critical to back up — they can be re-pulled.",
    "Update cadence: `ollama pull llama3.1:8b` re-pulls the latest tag. Plan a monthly model refresh.",
]:
    add_bullet(doc, item)

add_heading(doc, "17.8 Day-one validation checklist", level=2)
for item in [
    "✓ Ollama responds on Legion: `curl http://localhost:11434/api/tags`",
    "✓ Ollama reachable from Android over Tailscale: open the URL in phone browser, see 'Ollama is running'",
    "✓ llama3.1:8b warm: `ollama run llama3.1:8b 'hello'` returns within 1.5s",
    "✓ Embeddings work: `curl http://localhost:11434/api/embed -d '{\"model\":\"nomic-embed-text\",\"input\":\"test\"}'`",
    "✓ Tailscale Funnel URL works from outside the home Wi-Fi (try phone on LTE)",
    "✓ Supabase Edge Function can fetch the Funnel URL (test from Supabase function shell)",
    "If all six checks pass, the second brain is ready for the app to plug into.",
]:
    add_bullet(doc, item)


# Save — relative to this script's location
import os as _os
output_path = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "NotesApp_Architecture.docx")
doc.save(output_path)
print(f"OK: {output_path}")
